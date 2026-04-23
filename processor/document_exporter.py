"""
DocumentExporter – Exportiert Protokolle als DOCX, PDF und HTML.
Behörden-Stil mit optionaler Wappen-Einbindung.
"""

import os
import io
import base64
import logging
from typing import Optional, Dict, Any
from datetime import datetime

logger = logging.getLogger(__name__)


class DocumentExporter:
    """Exportiert Sitzungsprotokolle in verschiedene Formate."""

    def export_docx(
        self,
        protocol: Dict[str, Any],
        wappen_pfad: Optional[str] = None
    ) -> bytes:
        """
        Exportiert das Protokoll als DOCX (Word-Dokument) im Behördenstil.

        Args:
            protocol: Strukturiertes Protokoll
            wappen_pfad: Optionaler Pfad zum Wappen-Bild

        Returns:
            DOCX-Datei als Bytes
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE
        except ImportError:
            raise ImportError(
                "python-docx nicht installiert. Installieren mit: pip install python-docx"
            )

        doc = Document()

        # Seitenränder
        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        # Wappen einfügen
        if wappen_pfad and os.path.exists(wappen_pfad):
            try:
                doc.add_picture(wappen_pfad, width=Inches(1.5))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                logger.warning(f"Wappen konnte nicht eingefügt werden: {e}")

        # Titel
        title = doc.add_heading(protocol.get("titel", "Sitzungsprotokoll"), level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Behördenname
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(protocol.get("behoerdenname", ""))
        run.bold = True
        run.font.size = Pt(14)

        # Metadaten
        doc.add_paragraph()
        meta = doc.add_paragraph()
        meta.add_run("Datum: ").bold = True
        meta.add_run(str(protocol.get("datum", "")))
        meta.add_run("  |  Beginn: ").bold = True
        meta.add_run(f"{protocol.get('uhrzeit', '')} Uhr")
        meta.add_run("  |  Ort: ").bold = True
        meta.add_run(protocol.get("ort", ""))

        # Anwesende
        if protocol.get("sprecher"):
            p = doc.add_paragraph()
            p.add_run("Anwesende: ").bold = True
            p.add_run(", ".join(protocol["sprecher"]))

        doc.add_paragraph("─" * 60)

        # Tagesordnungspunkte
        doc.add_heading("Tagesordnungspunkte", level=2)
        for top in protocol.get("tagesordnungspunkte", []):
            doc.add_heading(
                f"TOP {top.get('nummer', '?')} – {top.get('titel', 'Ohne Titel')}",
                level=3
            )
            doc.add_paragraph(top.get("inhalt", ""))

        # Beschlüsse
        beschluesse = protocol.get("beschluesse", [])
        if beschluesse:
            doc.add_heading("Beschlüsse", level=2)
            for b in beschluesse:
                p = doc.add_paragraph()
                p.add_run(f"Beschluss {b.get('nummer', '?')}: ").bold = True
                p.add_run(b.get("text", ""))
                if b.get("abstimmung"):
                    p_abstimmung = doc.add_paragraph()
                    p_abstimmung.add_run("  Abstimmung: ").italic = True
                    p_abstimmung.add_run(b["abstimmung"]).italic = True

        # Zusammenfassung
        if protocol.get("zusammenfassung"):
            doc.add_heading("Zusammenfassung", level=2)
            doc.add_paragraph(protocol["zusammenfassung"])

        # Unterschriften
        doc.add_paragraph()
        doc.add_paragraph("─" * 60)
        doc.add_paragraph()

        sig_table = doc.add_table(rows=1, cols=2)
        sig_table.columns[0].width = Inches(3)
        sig_table.columns[1].width = Inches(3)
        row = sig_table.rows[0]
        row.cells[0].text = "_________________________\nProtokollführer/in"
        row.cells[1].text = "_________________________\nVorsitzende/r"

        # Fußzeile
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run(
            f"Automatisch generiert am {datetime.now().strftime('%d.%m.%Y')} "
            f"mit Sitzungsprotokoll-KI"
        )
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(128, 128, 128)

        # Bytes erstellen
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def export_pdf(
        self,
        protocol: Dict[str, Any],
        wappen_pfad: Optional[str] = None
    ) -> bytes:
        """
        Exportiert das Protokoll als PDF.

        Args:
            protocol: Strukturiertes Protokoll
            wappen_pfad: Optionaler Pfad zum Wappen-Bild

        Returns:
            PDF-Datei als Bytes
        """
        html_content = self.export_html(protocol, wappen_pfad)

        try:
            from weasyprint import HTML
            pdf_bytes = HTML(string=html_content).write_pdf()
            return pdf_bytes
        except ImportError:
            raise ImportError(
                "weasyprint nicht installiert. Installieren mit: pip install weasyprint"
            )
        except Exception as e:
            raise RuntimeError(f"PDF-Generierung fehlgeschlagen: {e}")

    def export_html(
        self,
        protocol: Dict[str, Any],
        wappen_pfad: Optional[str] = None
    ) -> str:
        """
        Exportiert das Protokoll als HTML mit eingebettetem CSS.

        Args:
            protocol: Strukturiertes Protokoll
            wappen_pfad: Optionaler Pfad zum Wappen-Bild

        Returns:
            HTML-String
        """
        # Wappen als Base64 einbetten
        wappen_html = ""
        if wappen_pfad and os.path.exists(wappen_pfad):
            try:
                with open(wappen_pfad, "rb") as f:
                    b64 = base64.b64encode(f.read()).decode("utf-8")
                ext = os.path.splitext(wappen_pfad)[1].lstrip(".")
                mime = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg"}.get(ext, "image/png")
                wappen_html = f'<img src="data:{mime};base64,{b64}" class="wappen" alt="Wappen">'
            except Exception as e:
                logger.warning(f"Wappen konnte nicht eingebettet werden: {e}")

        # Tagesordnungspunkte
        tops_html = ""
        for top in protocol.get("tagesordnungspunkte", []):
            tops_html += f"""
            <div class="top">
                <h3>TOP {top.get('nummer', '?')} – {self._escape_html(top.get('titel', ''))}</h3>
                <p>{self._escape_html(top.get('inhalt', ''))}</p>
            </div>
            """

        # Beschlüsse
        beschluesse_html = ""
        for b in protocol.get("beschluesse", []):
            abstimmung = f" <em>(Abstimmung: {self._escape_html(b['abstimmung'])})</em>" if b.get("abstimmung") else ""
            beschluesse_html += f"""
            <div class="beschluss">
                <strong>Beschluss {b.get('nummer', '?')}:</strong>
                {self._escape_html(b.get('text', ''))}{abstimmung}
            </div>
            """

        # Sprecher
        sprecher_str = ", ".join(protocol.get("sprecher", []))
        anwesende_html = f"<p><strong>Anwesende:</strong> {self._escape_html(sprecher_str)}</p>" if sprecher_str else ""

        html = f"""<!DOCTYPE html>
<html lang="de">
<head>
    <meta charset="UTF-8">
    <title>{self._escape_html(protocol.get('titel', 'Sitzungsprotokoll'))}</title>
    <style>
        body {{
            font-family: 'Arial', 'Helvetica', sans-serif;
            max-width: 800px;
            margin: 40px auto;
            padding: 20px;
            color: #333;
            line-height: 1.6;
        }}
        .wappen {{
            display: block;
            margin: 0 auto 20px;
            max-width: 150px;
            height: auto;
        }}
        h1 {{
            text-align: center;
            color: #1a365d;
            border-bottom: 3px solid #1a365d;
            padding-bottom: 10px;
        }}
        .behoerde {{
            text-align: center;
            font-size: 18px;
            font-weight: bold;
            margin-bottom: 5px;
        }}
        .meta {{
            text-align: center;
            color: #666;
            margin-bottom: 20px;
        }}
        .top {{
            margin: 20px 0;
            padding: 15px;
            background: #f8f9fa;
            border-left: 4px solid #1a365d;
            border-radius: 0 4px 4px 0;
        }}
        .top h3 {{
            color: #1a365d;
            margin-top: 0;
        }}
        .beschluss {{
            margin: 10px 0;
            padding: 10px 15px;
            background: #e8f5e9;
            border-left: 4px solid #2e7d32;
            border-radius: 0 4px 4px 0;
        }}
        .unterschrift {{
            display: flex;
            justify-content: space-between;
            margin-top: 60px;
        }}
        .unterschrift div {{
            text-align: center;
            width: 45%;
        }}
        .footer {{
            text-align: center;
            color: #999;
            font-size: 10px;
            margin-top: 40px;
            padding-top: 10px;
            border-top: 1px solid #ddd;
        }}
        @media print {{
            body {{ margin: 20px; }}
            .wappen {{ max-width: 100px; }}
        }}
    </style>
</head>
<body>
    {wappen_html}
    <h1>{self._escape_html(protocol.get('titel', 'Sitzungsprotokoll'))}</h1>
    <p class="behoerde">{self._escape_html(protocol.get('behoerdenname', ''))}</p>
    <p class="meta">
        Datum: {self._escape_html(str(protocol.get('datum', '')))} |
        Beginn: {self._escape_html(protocol.get('uhrzeit', ''))} Uhr |
        Ort: {self._escape_html(protocol.get('ort', ''))}
    </p>
    {anwesende_html}
    <hr>
    <h2>Tagesordnungspunkte</h2>
    {tops_html}
    {"<h2>Beschlüsse</h2>" + beschluesse_html if beschluesse_html else ""}
    {"<h2>Zusammenfassung</h2><p>" + self._escape_html(protocol.get('zusammenfassung', '')) + "</p>" if protocol.get('zusammenfassung') else ""}
    <div class="unterschrift">
        <div>_________________________<br>Protokollführer/in</div>
        <div>_________________________<br>Vorsitzende/r</div>
    </div>
    <p class="footer">Automatisch generiert am {datetime.now().strftime('%d.%m.%Y')} mit Sitzungsprotokoll-KI</p>
</body>
</html>"""
        return html

    def _escape_html(self, text: str) -> str:
        """Escapes HTML-Sonderzeichen."""
        if not text:
            return ""
        return (
            text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;")
        )
